# -*- coding: utf-8 -*-
import sys
from docx import Document

# UTF-8 출력 설정
sys.stdout.reconfigure(encoding='utf-8')

doc_path = r'C:\Users\dbsrj\Downloads\IsItLegal_기획안_v1.0.docx'
doc = Document(doc_path)

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# 테이블 내용도 추출
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                print(cell.text)
